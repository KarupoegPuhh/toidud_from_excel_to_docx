
import docx
from docx import Document
from docx import table
from docx.shared import Cm
from openpyxl import load_workbook
import sys
import os
import datetime


def copy_tbl(y1, x1, y2, x2, sheet):
    tükk = []
    for i in range(x1, x2 + 1, 1):
        rida = []
        for j in range(y1, y2 + 1, 1):
            rida.append(sheet.cell(i, j).value)
        tükk.append(rida)
    return tükk

def get_para_data(output_dok, paragraph, style, append_last_run_with=""):

    paragraaf_loplik = output_dok.add_paragraph(style=style)
    for run in paragraph.runs:
        output_run = paragraaf_loplik.add_run(run.text)
        output_run.bold = run.bold
        output_run.italic = run.italic
        output_run.underline = run.underline
        output_run.font.size = run.font.size
        #print("fnt", run.font.size)
        output_run.font.name = run.font.name
        #output_run.style.name = run.style.name

    if append_last_run_with != "":
        paragraaf_loplik.runs[-1].text += append_last_run_with
        #print("OUTPUT", paragraaf_loplik.text)

    #paragrafi alignment data
    paragraaf_loplik.paragraph_format.alignment = paragraph.paragraph_format.alignment
    paragraaf_loplik.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
    paragraaf_loplik.paragraph_format.space_after = 0
    #print("spc",paragraph.paragraph_format.space_after)

def datetime_to_est(k):
    kuud = ["jaanuar",  "veebruar",  "märts",  "aprill",  "mai",  "juuni",  "juuli",  "august",  "september",  "oktoober",  "november",  "detsember"]
    kuu = kuud[int(k.strftime("%m"))-1]
    päev = int(k.strftime("%d"))-1
    return str(päev) +". "+ kuu +" "+ k.strftime("%Y")


if len(sys.argv) == 1:
    print("Programm käivitati ilma parameetrita (toitude tabel tuleb lohistada main.exe peale)")
    exit(1)
print("toidutabel: "+sys.argv[1])
#toitude_tabel_faili_nimi = input("sisesta faili nimi kus on toidud (nt: 30 august - 3 september 2021a.xlsx) ")
toitude_tabel_faili_tee = sys.argv[1]
wb = load_workbook(toitude_tabel_faili_tee)
toitude_tabel_lehe_nimi = os.path.basename(toitude_tabel_faili_tee).replace(".xlsx", "")
sheet = wb[toitude_tabel_lehe_nimi]

try:
    wb2 = load_workbook("lastehoiud ja kuupäevad.xlsx", data_only=True)
except:
    print("ei saanud kätte: lastehoiud ja kuupäevad.xlsx")
    exit(1)

hoiud = wb2.active
#print(hoiud)
#hoiud = wb["Leht1"]
hoide = 0
i = 2
while hoiud.cell(i,1).value != None:
    hoide +=1
    i+=1
#print(hoide)

dok = Document("data.saatelehed.docx")
dok2 = Document()
#margins
sections = dok2.sections
for section in sections:
    section.top_margin = dok.sections[0].top_margin
    section.bottom_margin = dok.sections[0].bottom_margin
    section.left_margin = dok.sections[0].left_margin
    section.right_margin = dok.sections[0].right_margin

#s = []
#stiil =[]

for hoid in range(hoide):
    mitmes_saateleht_hoiule = 0
    for päev in range(5):
        i = 0
        kuupäev = hoiud.cell(row=hoid + 2, column=päev + 2).value
        menüü = copy_tbl(päev * 3 + 1, 5, päev * 3 + 3, 13, sheet)
        #print(menüü[0][0])
        if kuupäev == None:
            print(str(päev+1)+". kuupäev failis 'lastehoiud ja kuupäevad' puudub, saatelehte ei genereeri")
        elif menüü[0][0] == None:
            print("menüü "+str(päev+1)+". päeva all pole toite, saatelehte ei genereeri")
        else:
            for r in dok.paragraphs:
                #for run in r.runs:
                #    s.append(run.text)
                #stiil.append(r.style)
                #print(r.style)
                #dok2.add_paragraph(s[i], stiil[i])

                if i==3:
                    table1 = dok2.add_table(1, 2)
                    #table1.style = "TableGrid"
                    hdr_cells = table1.rows[0].cells
                    hdr_cells[0].text = 'Toidu väljastamise kuupäev:'
                    #print(datetime_to_est(kuupäev))
                    hdr_cells[1].text = datetime_to_est(kuupäev)
                    hdr_cells[1].paragraphs[0].runs[0].bold = True
                    row_cells = table1.add_row().cells
                    row_cells[0].text = "Toidu väljastamise kellaaeg: "
                    row_cells[1].text = "Toit säilib kuni (kellaaeg): "

                if i==6:
                    table = dok2.add_table(1,5)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Jrk. nr'
                    hdr_cells[1].text = 'Toidu nimetus'
                    hdr_cells[2].text = 'Portsu kaal  keskmiselt'
                    hdr_cells[3].text = 'Kcal keskmiselt'
                    hdr_cells[4].text = 'Portsjonite arv'
                    hdr_cells[0].paragraphs[0].runs[0].bold = True
                    hdr_cells[1].paragraphs[0].runs[0].bold = True
                    hdr_cells[2].paragraphs[0].runs[0].bold = True
                    hdr_cells[3].paragraphs[0].runs[0].bold = True
                    hdr_cells[4].paragraphs[0].runs[0].bold = True
                    j=0
                    for toit, kaal, kcal in menüü:
                        j+=1
                        if toit != "" and toit != None and kcal != None and kaal != None:
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(j)
                            row_cells[1].text = toit
                            row_cells[2].text = str(kaal)+" g"
                            row_cells[3].text = str(round(kcal,2))+" kcal"
                    #veergude laiused
                    for cell in table.columns[0].cells:
                        cell.width = Cm(1.19)
                    for cell in table.columns[1].cells:
                        cell.width = Cm(8)
                    for cell in table.columns[2].cells:
                        cell.width = Cm(3)
                    for cell in table.columns[3].cells:
                        cell.width = Cm(2.75)
                    for cell in table.columns[3].cells:
                        cell.width = Cm(3)
                if i == 1:
                    #print(r.text)
                    get_para_data(dok2, r, r.style, str(hoiud.cell(row=hoid + 2, column=1).value))
                else:
                    get_para_data(dok2, r, r.style)
                if i == 2:
                    u = dok2.add_paragraph()
                    u.paragraph_format.space_after = 0
                #print(i)
                #print(r.text)
                if i==9:
                    print("Genereeritud saateleht kuupäevaga", datetime_to_est(kuupäev)+".", "Saaja:", str(hoiud.cell(row=hoid + 2, column=1).value))
                i+=1
            if mitmes_saateleht_hoiule%2 != 0:
                dok2.add_page_break()
            else:
                a = dok2.add_paragraph()
                a.paragraph_format.space_after = 0
                b = dok2.add_paragraph()
                b.paragraph_format.space_after = 0
            mitmes_saateleht_hoiule += 1
    dok2.add_page_break()
#print('\n'.join(s))
dok2.save('genereeritud_saatelehed.docx')